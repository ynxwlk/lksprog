#coding=utf-8

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
#打开文档
document = Document()

#设置中文字体
paragraph = document.add_paragraph()
run = paragraph.add_run(u'曲靖市规划局麒麟分局\n')
run.font.size = Pt(24)
run.font.name=u'宋体'


run = paragraph.add_run(u'建设工程规划许可证初审表')
run.font.size = Pt(24)
run.font.name=u'宋体'


#增加表格
table = document.add_table(rows=12, cols=5,style = 'Table Grid')

for num in range(5):
    table.cell(num,0).merge(table.cell(num,1))
    if (num < 3):
        table.cell(num,2).merge(table.cell(num,4))

table.cell(3,0).merge(table.cell(4,1))  #办证内容
table.cell(3,2).merge(table.cell(4,2))  #建设工程规划许可证

for num in range(6):
    table.cell(num + 5, 1 ).merge(table.cell(num + 5 , 2 ))
    table.cell(num + 5, 3 ).merge(table.cell(num + 5 , 4 ))

table.cell(5,0).merge(table.cell(10,0))
table.cell(11, 1 ).merge(table.cell(11 , 4 ))
table.cell(5,0).width = Inches(0.5)
table.cell(11,0).width = Inches(1)

#保存文件
document.save(u'建设工程规划许可证初审表.docx')